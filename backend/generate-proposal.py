#!/usr/bin/env python3
"""Generate professional DOCX proposal for Kedai Ngupi-Ngupi WhatsApp AI Chatbot."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Helper functions ──
def add_blank(n=1):
    for _ in range(n):
        doc.add_paragraph('')

def add_title_text(text, size=28, bold=True, color=RGBColor(0x1B, 0x3A, 0x4B)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)
        run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 + level * 0.75)
    return p

def add_numbered(text):
    return doc.add_paragraph(text, style='List Number')

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
add_blank(4)
add_title_text('PROPOSAL IMPLEMENTASI', size=26)
add_title_text('WHATSAPP AI CHATBOT', size=30)
add_blank(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('KEDAI NGUPI-NGUPI')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x8B, 0x5C, 0x2A)
run.font.name = 'Calibri'

add_blank(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Purwakarta, Jawa Barat')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_blank(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Disusun oleh:')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Muhammad Rasyid Ridho')
run.bold = True
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Developer WhatsApp AI Chatbot')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_blank(1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('April 2026')
run.bold = True
run.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
add_heading_styled('DAFTAR ISI', level=1)
toc_items = [
    '1. Pendahuluan',
    '   1.1 Latar Belakang',
    '   1.2 Rumusan Masalah',
    '   1.3 Tujuan Implementasi',
    '2. Gambaran Umum Solusi',
    '3. Visi dan Misi',
    '4. Ruang Lingkup Fitur',
    '5. Analisis Pasar',
    '6. Analisis Layanan',
    '7. Alur Operasional Sistem',
    '8. Arsitektur Teknis',
    '9. Manfaat Implementasi',
    '10. Estimasi Timeline',
    '11. Estimasi Anggaran',
    '12. Deliverables',
    '13. Risiko dan Mitigasi',
    '14. Penutup',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. PENDAHULUAN
# ══════════════════════════════════════════════════════════════
add_heading_styled('1. Pendahuluan', level=1)

add_heading_styled('1.1 Latar Belakang', level=2)
add_para(
    'Perkembangan industri makanan dan minuman (F&B) saat ini berjalan seiring dengan '
    'meningkatnya kebutuhan konsumen terhadap layanan yang cepat, praktis, dan mudah diakses. '
    'Perubahan perilaku pelanggan mendorong pelaku usaha untuk tidak hanya menghadirkan produk '
    'yang berkualitas, tetapi juga sistem pelayanan yang efisien dan responsif.'
)
add_para(
    'Dalam operasional sehari-hari, proses pemesanan manual melalui chat sering menimbulkan '
    'beberapa kendala, seperti keterlambatan respons, potensi kesalahan pencatatan pesanan, '
    'keterbatasan pelayanan di jam sibuk, serta tingginya ketergantungan pada ketersediaan staf. '
    'Selain itu, penggunaan platform pesan antar pihak ketiga juga menimbulkan beban biaya komisi '
    'yang cukup besar bagi pelaku usaha. Potongan komisi merchant yang dapat mencapai sekitar 20% '
    'per transaksi berdampak langsung pada margin keuntungan, sehingga banyak merchant akhirnya '
    'menaikkan harga menu di platform agar tetap memperoleh margin yang wajar.'
)
add_para(
    'Kondisi ini membuka peluang bagi Kedai Ngupi-Ngupi untuk membangun jalur pemesanan mandiri '
    'yang lebih efisien melalui WhatsApp, yaitu aplikasi pesan instan yang paling banyak digunakan '
    'di Indonesia. Dengan implementasi WhatsApp AI Chatbot, kedai dapat menghadirkan layanan '
    'pemesanan yang responsif, otomatis, dan terintegrasi dengan sistem pembayaran digital, tanpa '
    'harus menambah beban kerja staf secara berlebihan.'
)
add_para(
    'Sistem ini dirancang sebagai asisten virtual yang dapat melayani pelanggan untuk berbagai '
    'kebutuhan, mulai dari pemesanan dine-in, pickup, delivery melalui layanan Go Ngupi, hingga '
    'reservasi meja. Dengan pendekatan ini, Kedai Ngupi-Ngupi tidak hanya memperoleh efisiensi '
    'operasional, tetapi juga membangun ekosistem penjualan mandiri yang lebih sehat secara finansial '
    'dan lebih dekat dengan pelanggan.'
)

add_heading_styled('1.2 Rumusan Masalah', level=2)
add_para('Berdasarkan kondisi operasional yang ada, beberapa permasalahan yang ingin dijawab melalui implementasi sistem ini adalah:')
problems = [
    'Bagaimana cara meningkatkan kecepatan dan konsistensi respons terhadap pelanggan, terutama pada jam sibuk?',
    'Bagaimana cara mengurangi ketergantungan pada proses pencatatan manual yang rawan kesalahan?',
    'Bagaimana cara membangun kanal pemesanan mandiri yang dapat menekan ketergantungan pada platform pihak ketiga beserta beban komisinya?',
    'Bagaimana cara menghadirkan proses pemesanan dan pembayaran yang lebih praktis dan terintegrasi bagi pelanggan?',
    'Bagaimana cara mengelola data pelanggan dan transaksi secara lebih terstruktur untuk kebutuhan bisnis ke depan?',
]
for i, prob in enumerate(problems, 1):
    add_numbered(prob)

add_heading_styled('1.3 Tujuan Implementasi', level=2)
add_para('Tujuan dari implementasi WhatsApp AI Chatbot ini adalah:')
goals = [
    'Membangun sistem pemesanan mandiri berbasis WhatsApp yang dapat diakses oleh pelanggan kapan saja.',
    'Meningkatkan efisiensi pelayanan pelanggan melalui otomasi alur percakapan dan pemesanan.',
    'Mengurangi beban operasional staf dalam menangani pertanyaan dan pesanan yang bersifat rutin dan berulang.',
    'Mendukung proses pembayaran digital yang lebih cepat, aman, dan terintegrasi langsung dalam percakapan WhatsApp.',
    'Membangun database pelanggan dan histori transaksi sebagai aset data untuk pengembangan bisnis.',
    'Menciptakan transparansi harga bagi pelanggan, di mana harga delivery dapat disamakan dengan harga dine-in.',
]
for g in goals:
    add_numbered(g)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. GAMBARAN UMUM SOLUSI
# ══════════════════════════════════════════════════════════════
add_heading_styled('2. Gambaran Umum Solusi', level=1)
add_para(
    'WhatsApp AI Chatbot adalah sistem asisten virtual berbasis kecerdasan buatan yang dirancang '
    'untuk membantu Kedai Ngupi-Ngupi dalam menangani interaksi pelanggan secara otomatis melalui '
    'WhatsApp. Sistem ini bekerja dengan alur percakapan yang terstruktur sehingga pelanggan dapat '
    'melakukan pemesanan, memperoleh informasi menu, memilih metode layanan, hingga menerima '
    'instruksi pembayaran tanpa harus menunggu respons manual dari staf.'
)
add_para(
    'Dalam implementasinya, chatbot tidak menggantikan peran manusia sepenuhnya. Sistem ini '
    'berfungsi sebagai lapisan awal pelayanan yang menangani kebutuhan rutin dan berulang. Untuk '
    'kebutuhan khusus seperti komplain, kendala teknis, atau permintaan yang memerlukan keputusan '
    'manual, sistem tetap menyediakan mekanisme eskalasi ke admin manusia.'
)
add_para(
    'Secara teknis, sistem ini terdiri dari beberapa komponen utama yang saling terintegrasi:'
)
components = [
    'AI Engine: memproses pesan pelanggan, memahami konteks, dan menghasilkan respons yang relevan.',
    'Order Management: mengelola alur pemesanan dari awal hingga selesai, termasuk konfirmasi dan status tracking.',
    'Payment Integration: terhubung dengan payment gateway untuk mendukung pembayaran digital secara real-time.',
    'Customer Database: menyimpan profil pelanggan, histori transaksi, dan preferensi untuk personalisasi layanan.',
    'Admin Dashboard: menyediakan antarmuka bagi staf untuk monitoring, takeover, dan pengelolaan operasional.',
]
for c in components:
    add_bullet(c)

add_para(
    'Dengan pendekatan ini, Kedai Ngupi-Ngupi dapat meningkatkan kualitas layanan, mempercepat '
    'proses transaksi, dan menjaga konsistensi operasional, terutama pada jam-jam sibuk di mana '
    'volume interaksi pelanggan meningkat.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. VISI DAN MISI
# ══════════════════════════════════════════════════════════════
add_heading_styled('3. Visi dan Misi', level=1)

add_heading_styled('3.1 Visi', level=2)
add_para(
    'Menjadi kedai kopi modern yang mandiri, efisien, dan berorientasi pada pengalaman pelanggan '
    'melalui pemanfaatan teknologi digital yang relevan dan terintegrasi.',
    italic=True
)

add_heading_styled('3.2 Misi', level=2)
missions = [
    'Menghadirkan layanan pemesanan yang cepat, praktis, dan mudah diakses melalui WhatsApp sebagai kanal utama.',
    'Mengembangkan ekosistem delivery mandiri Go Ngupi yang lebih efisien secara biaya dibandingkan platform pihak ketiga.',
    'Mendorong transparansi harga bagi pelanggan melalui kanal penjualan langsung tanpa beban komisi eksternal.',
    'Mengintegrasikan proses pemesanan, pembayaran, dan pencatatan transaksi dalam satu alur kerja yang rapi dan terukur.',
    'Membantu staf fokus pada kualitas produk dan pelayanan langsung di outlet dengan mengurangi pekerjaan administratif berulang.',
]
for i, m in enumerate(missions, 1):
    add_numbered(m)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. RUANG LINGKUP FITUR
# ══════════════════════════════════════════════════════════════
add_heading_styled('4. Ruang Lingkup Fitur', level=1)
add_para('Sistem WhatsApp AI Chatbot yang diusulkan mencakup fitur-fitur berikut:')

features = [
    ('4.1 Pemesanan Dine-in',
     'Pelanggan dapat melakukan pemesanan untuk konsumsi di tempat melalui alur yang terstruktur. '
     'Sistem mendukung identifikasi meja melalui QR code yang tersedia di setiap meja, sehingga '
     'pesanan dapat langsung terhubung dengan posisi pelanggan di outlet.'),
    ('4.2 Pemesanan Pickup',
     'Pelanggan dapat memesan terlebih dahulu melalui WhatsApp dan mengambil pesanan di outlet '
     'setelah selesai disiapkan. Sistem akan memberikan notifikasi kesiapan pesanan.'),
    ('4.3 Pemesanan Delivery (Go Ngupi)',
     'Pelanggan dapat melakukan pemesanan antar melalui armada delivery internal kedai. Sistem '
     'mendukung validasi lokasi pengiriman, perhitungan ongkos kirim berdasarkan jarak, dan '
     'koordinasi dengan kurir internal.'),
    ('4.4 Reservasi Meja',
     'Sistem dapat mencatat kebutuhan reservasi pelanggan secara terstruktur, meliputi tanggal, '
     'waktu, jumlah tamu, dan nama pemesan.'),
    ('4.5 Katalog Menu Interaktif',
     'Chatbot dapat menampilkan kategori menu, daftar item, harga, ketersediaan produk, opsi '
     'varian, dan deskripsi singkat. Menu diperbarui secara otomatis dari sistem POS.'),
    ('4.6 Integrasi Pembayaran Digital',
     'Sistem terhubung dengan payment gateway (Doku) untuk mendukung pembayaran QRIS, '
     'e-wallet, dan metode nontunai lainnya. Invoice dan QR pembayaran dikirim langsung ke '
     'ruang obrolan WhatsApp pelanggan.'),
    ('4.7 Verifikasi Pembayaran Otomatis',
     'Status pembayaran diverifikasi secara otomatis melalui callback dari payment gateway. '
     'Pesanan yang sudah lunas langsung diperbarui statusnya dan diteruskan ke tim produksi.'),
    ('4.8 Human Takeover',
     'Percakapan dapat dialihkan ke admin manusia apabila pelanggan menyampaikan komplain, '
     'pertanyaan kompleks, atau kebutuhan di luar skenario otomatis.'),
    ('4.9 Database Pelanggan',
     'Sistem menyimpan profil pelanggan, riwayat transaksi, preferensi menu, dan pola pembelian '
     'untuk mendukung personalisasi layanan, program loyalitas, dan kegiatan promosi.'),
]
for title, desc in features:
    add_heading_styled(title, level=2)
    add_para(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. ANALISIS PASAR
# ══════════════════════════════════════════════════════════════
add_heading_styled('5. Analisis Pasar', level=1)

add_heading_styled('5.1 Target Pasar', level=2)
add_para('Target utama dari implementasi sistem ini meliputi:')
targets = [
    'Pelanggan loyal Kedai Ngupi-Ngupi yang sudah terbiasa memesan melalui WhatsApp.',
    'Pekerja kantoran dan profesional muda yang membutuhkan layanan cepat dan praktis.',
    'Mahasiswa dan komunitas lokal di area Purwakarta dan sekitarnya.',
    'Konsumen dengan mobilitas tinggi yang mengutamakan kemudahan transaksi digital.',
    'Pelanggan yang sensitif terhadap perbedaan harga antara dine-in dan delivery di platform lain.',
]
for t in targets:
    add_bullet(t)

add_heading_styled('5.2 Kondisi Pasar', level=2)
add_para(
    'Saat ini, mayoritas pelaku usaha F&B di Indonesia masih mengandalkan platform agregator '
    'pihak ketiga untuk layanan delivery. Model ini membantu dari sisi jangkauan dan kemudahan, '
    'tetapi memberikan tekanan pada margin melalui potongan komisi merchant yang dapat mencapai '
    'sekitar 20% per transaksi. Akibatnya, banyak merchant menaikkan harga menu di platform, '
    'sehingga pelanggan membayar lebih mahal untuk produk yang sama.'
)

add_heading_styled('5.3 Celah Pasar', level=2)
add_para(
    'Pelanggan mulai menyadari adanya perbedaan harga antara pembelian langsung dan pembelian '
    'melalui platform. Di titik ini, kanal pemesanan mandiri memiliki daya tarik yang kuat karena '
    'dapat menawarkan harga yang lebih wajar dengan proses yang tetap praktis. Tren conversational '
    'commerce juga menunjukkan bahwa konsumen semakin nyaman bertransaksi melalui aplikasi pesan '
    'instan yang sudah mereka gunakan sehari-hari.'
)

add_heading_styled('5.4 Keunggulan Kompetitif', level=2)
add_para('Keunggulan utama Kedai Ngupi-Ngupi melalui sistem ini:')
usps = [
    'Harga delivery sama dengan harga dine-in karena tidak ada beban komisi platform.',
    'Pemesanan langsung via WhatsApp tanpa perlu mengunduh aplikasi tambahan.',
    'Interaksi lebih personal dan responsif melalui AI yang memahami konteks percakapan.',
    'Data pelanggan dimiliki langsung oleh kedai untuk pengembangan bisnis.',
    'Ekosistem delivery internal (Go Ngupi) yang lebih fleksibel dan terkendali.',
]
for u in usps:
    add_bullet(u)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. ANALISIS LAYANAN
# ══════════════════════════════════════════════════════════════
add_heading_styled('6. Analisis Layanan', level=1)

add_heading_styled('6.1 Deskripsi Layanan Inti', level=2)
add_para(
    'WhatsApp AI Chatbot berfungsi sebagai asisten virtual yang menangani interaksi pelanggan '
    'secara otomatis. Layanan inti yang diberikan meliputi penerimaan pesanan, pengelompokan '
    'jenis layanan, konfirmasi rincian order, pengiriman instruksi pembayaran, pencatatan '
    'reservasi, pemberian informasi menu, dan pengalihan ke admin jika dibutuhkan.'
)

add_heading_styled('6.2 Sistem Logistik On-Demand', level=2)
add_para(
    'Layanan delivery Go Ngupi dikelola secara on-demand dengan pendekatan point-to-point. '
    'Chatbot memproses lokasi pengiriman pelanggan, menghitung jarak dan ongkos kirim secara '
    'otomatis, lalu menyinkronkan data pesanan dengan kesiapan kurir internal. Pendekatan ini '
    'memungkinkan proses pengantaran yang lebih cepat dan efisien.'
)

add_heading_styled('6.3 Integrasi Pembayaran Tersentralisasi', level=2)
add_para(
    'Sistem terintegrasi langsung dengan payment gateway (Doku) untuk mendukung penerbitan '
    'invoice otomatis, pembayaran QRIS, e-wallet, dan metode nontunai lainnya. Seluruh proses '
    'pembayaran berlangsung di dalam ruang obrolan WhatsApp tanpa pelanggan harus berpindah '
    'aplikasi.'
)

add_heading_styled('6.4 Skalabilitas Operasional', level=2)
add_para(
    'Sistem chatbot dapat menangani banyak percakapan secara bersamaan tanpa penurunan kualitas '
    'respons. Pada jam sibuk, di mana volume interaksi pelanggan meningkat, sistem tetap dapat '
    'memberikan pelayanan yang konsisten. Kapasitas ini sulit dicapai jika hanya mengandalkan '
    'respons manual dari staf.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. ALUR OPERASIONAL SISTEM
# ══════════════════════════════════════════════════════════════
add_heading_styled('7. Alur Operasional Sistem', level=1)
add_para(
    'Alur operasional sistem dirancang agar sederhana bagi pelanggan dan efisien bagi tim '
    'operasional kedai. Berikut adalah tahapan utama:'
)

steps = [
    ('7.1 Inisiasi Percakapan',
     'Pelanggan menghubungi nomor WhatsApp resmi Kedai Ngupi-Ngupi. Sistem memberikan sapaan '
     'awal dan menampilkan pilihan layanan utama: Dine-in, Pickup, Delivery (Go Ngupi), '
     'Reservasi, atau Bantuan Admin.'),
    ('7.2 Pemilihan Layanan dan Input Pesanan',
     'Setelah memilih layanan, pelanggan diarahkan ke alur yang sesuai. Chatbot menampilkan '
     'menu atau menerima item yang dipesan, lalu memandu pelanggan untuk melengkapi informasi '
     'yang dibutuhkan seperti jumlah item, varian, nama pelanggan, dan lokasi pengiriman '
     'bila diperlukan.'),
    ('7.3 Konfirmasi Ringkasan Pesanan',
     'Setelah semua data terkumpul, chatbot menampilkan ringkasan pesanan yang mencakup nama '
     'pelanggan, item dan jumlah, harga per item, subtotal, ongkos kirim (jika delivery), '
     'dan total pembayaran. Pelanggan diminta memberikan persetujuan sebelum melanjutkan.'),
    ('7.4 Pembayaran',
     'Untuk transaksi nontunai, sistem mengirimkan QR pembayaran atau invoice melalui payment '
     'gateway yang terintegrasi. Untuk skenario tertentu, sistem juga mendukung metode COD '
     '(Cash on Delivery) sesuai kebijakan operasional kedai.'),
    ('7.5 Verifikasi dan Pemrosesan',
     'Setelah pembayaran terkonfirmasi, status order diperbarui otomatis menjadi siap diproses. '
     'Pesanan diteruskan ke tim operasional atau dashboard internal untuk segera disiapkan.'),
    ('7.6 Penyelesaian Pesanan',
     'Barista atau dapur menerima data order yang sudah valid. Untuk pickup dan dine-in, '
     'pelanggan diberi notifikasi bahwa pesanan sedang disiapkan. Untuk delivery, order '
     'diteruskan ke alur pengantaran Go Ngupi.'),
    ('7.7 Eskalasi ke Admin',
     'Jika chatbot tidak dapat memahami kebutuhan pelanggan, menemukan kasus khusus, atau '
     'menerima komplain, percakapan dialihkan ke admin manusia untuk penanganan lanjutan.'),
]
for title, desc in steps:
    add_heading_styled(title, level=2)
    add_para(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. ARSITEKTUR TEKNIS
# ══════════════════════════════════════════════════════════════
add_heading_styled('8. Arsitektur Teknis', level=1)
add_para('Agar sistem berjalan dengan baik, diperlukan beberapa komponen teknis sebagai berikut:')

# Table
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['No.', 'Komponen', 'Keterangan']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

rows_data = [
    ('1', 'Nomor WhatsApp Bisnis', 'Nomor resmi yang digunakan sebagai kanal komunikasi pelanggan.'),
    ('2', 'Server / Cloud Hosting', 'Untuk menjalankan engine chatbot, logika order, database, dan integrasi.'),
    ('3', 'AI Engine', 'Memproses pesan pelanggan, mengelola alur percakapan, dan menjalankan logika pemesanan.'),
    ('4', 'Database', 'Menyimpan histori order, profil pelanggan, data reservasi, dan konfigurasi menu.'),
    ('5', 'Payment Gateway', 'Doku atau provider sejenis untuk mendukung QRIS, e-wallet, dan pembayaran digital.'),
    ('6', 'Admin Dashboard', 'Monitoring order, takeover manual, pelacakan status transaksi, dan laporan.'),
    ('7', 'Integrasi POS', 'Sinkronisasi data menu, harga, dan ketersediaan produk dengan sistem POS (Pawoon).'),
]
for i, (no, comp, desc) in enumerate(rows_data):
    row = table.rows[i + 1]
    row.cells[0].text = no
    row.cells[1].text = comp
    row.cells[2].text = desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

add_blank(1)
add_para(
    'Arsitektur sistem dirancang modular sehingga setiap komponen dapat diperbarui atau '
    'ditingkatkan secara independen tanpa mengganggu komponen lainnya.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. MANFAAT IMPLEMENTASI
# ══════════════════════════════════════════════════════════════
add_heading_styled('9. Manfaat Implementasi', level=1)

benefits = [
    ('Peningkatan Kecepatan Pelayanan',
     'Pelanggan memperoleh respons awal dalam hitungan detik tanpa harus menunggu admin membalas secara manual.'),
    ('Efisiensi Operasional',
     'Staf dapat lebih fokus pada produksi dan pelayanan langsung di outlet karena beban respons rutin ditangani oleh sistem.'),
    ('Pengurangan Kesalahan Pencatatan',
     'Data pesanan dicatat melalui alur yang terstruktur sehingga mengurangi potensi human error dalam pencatatan item, nominal, dan detail varian.'),
    ('Peningkatan Kapasitas Pelayanan',
     'Sistem dapat menangani banyak percakapan sekaligus, terutama pada jam sibuk, tanpa perlu menambah jumlah staf frontliner.'),
    ('Pembangunan Kanal Penjualan Mandiri',
     'Kedai memiliki jalur transaksi langsung dengan pelanggan tanpa sepenuhnya bergantung pada platform pihak ketiga dan beban komisinya.'),
    ('Pengelolaan Data Pelanggan',
     'Histori transaksi dan preferensi pelanggan dapat dimanfaatkan untuk analisis penjualan, program loyalitas, dan kegiatan promosi yang lebih terarah.'),
    ('Peningkatan Pengalaman Pelanggan',
     'Pelanggan mendapatkan proses pemesanan yang cepat, jelas, dan konsisten dalam satu aplikasi yang sudah mereka gunakan setiap hari.'),
]
for i, (title, desc) in enumerate(benefits, 1):
    add_heading_styled(f'9.{i} {title}', level=2)
    add_para(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. ESTIMASI TIMELINE
# ══════════════════════════════════════════════════════════════
add_heading_styled('10. Estimasi Timeline Implementasi', level=1)
add_para('Estimasi pelaksanaan proyek versi awal adalah sebagai berikut:')

# Timeline table
tt = doc.add_table(rows=5, cols=3)
tt.style = 'Light Grid Accent 1'
tt.alignment = WD_TABLE_ALIGNMENT.CENTER

tt_headers = ['Periode', 'Fase', 'Aktivitas']
for i, h in enumerate(tt_headers):
    cell = tt.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

tt_data = [
    ('Minggu 1', 'Discovery & Planning',
     'Pengumpulan kebutuhan, finalisasi flow percakapan, mapping menu, penentuan payment gateway, penentuan skenario operasional.'),
    ('Minggu 2', 'Development',
     'Pengembangan chatbot flow, integrasi database, integrasi payment gateway, penyusunan logika order dan status transaksi.'),
    ('Minggu 3', 'Testing & QA',
     'Pengujian end-to-end, simulasi pemesanan untuk semua skenario layanan (dine-in, pickup, delivery, reservasi), revisi berdasarkan hasil testing.'),
    ('Minggu 4', 'Deployment & Go-Live',
     'Deployment sistem, pendampingan go-live, monitoring awal, penyesuaian minor pasca implementasi.'),
]
for i, (period, phase, activities) in enumerate(tt_data):
    row = tt.rows[i + 1]
    row.cells[0].text = period
    row.cells[1].text = phase
    row.cells[2].text = activities
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

add_blank(1)
add_para(
    'Estimasi total durasi implementasi: 2-4 minggu, menyesuaikan kompleksitas kebutuhan '
    'dan kesiapan operasional.',
    bold=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. ESTIMASI ANGGARAN
# ══════════════════════════════════════════════════════════════
add_heading_styled('11. Estimasi Anggaran', level=1)
add_para('Estimasi anggaran implementasi dapat dibagi menjadi beberapa komponen utama:')

# Budget table
bt = doc.add_table(rows=6, cols=4)
bt.style = 'Light Grid Accent 1'
bt.alignment = WD_TABLE_ALIGNMENT.CENTER

bt_headers = ['No.', 'Komponen', 'Jenis Biaya', 'Estimasi']
for i, h in enumerate(bt_headers):
    cell = bt.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

bt_data = [
    ('1', 'Setup & Development',
     'Satu kali (one-time)',
     'Rp8.000.000 - Rp15.000.000'),
    ('2', 'Server & Infrastruktur',
     'Bulanan',
     'Rp300.000 - Rp1.000.000/bln'),
    ('3', 'WhatsApp API & Messaging',
     'Bulanan',
     'Rp300.000 - Rp2.000.000/bln'),
    ('4', 'Maintenance & Support',
     'Bulanan',
     'Rp500.000 - Rp2.500.000/bln'),
    ('5', 'Pengembangan Lanjutan',
     'Opsional',
     'Menyesuaikan scope'),
]
for i, (no, comp, jenis, est) in enumerate(bt_data):
    row = bt.rows[i + 1]
    row.cells[0].text = no
    row.cells[1].text = comp
    row.cells[2].text = jenis
    row.cells[3].text = est
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

add_blank(1)
add_para('Rincian per komponen:')

add_heading_styled('11.1 Setup dan Development', level=2)
add_para('Meliputi perancangan sistem, pengembangan flow chatbot, integrasi database, integrasi payment gateway, pengujian, dan deployment awal.')

add_heading_styled('11.2 Server dan Infrastruktur', level=2)
add_para('Meliputi cloud hosting/VPS, database, monitoring, dan backup.')

add_heading_styled('11.3 WhatsApp API dan Messaging Cost', level=2)
add_para('Meliputi biaya provider WhatsApp API dan/atau biaya percakapan sesuai volume penggunaan.')

add_heading_styled('11.4 Maintenance dan Support', level=2)
add_para('Meliputi monitoring sistem, perbaikan bug, penyesuaian flow/menu, dan dukungan teknis.')

add_heading_styled('11.5 Pengembangan Lanjutan (Opsional)', level=2)
add_para('Jika ke depan diperlukan fitur tambahan seperti program loyalitas, auto broadcast promo, dashboard analitik, integrasi POS lebih dalam, atau dukungan multi-cabang.')

add_para('Catatan: nominal di atas merupakan estimasi awal dan dapat berubah sesuai ruang lingkup final proyek.', italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. DELIVERABLES
# ══════════════════════════════════════════════════════════════
add_heading_styled('12. Deliverables', level=1)
add_para('Pada akhir implementasi, output yang diterima oleh Kedai Ngupi-Ngupi meliputi:')

deliverables = [
    'WhatsApp AI Chatbot aktif dan siap digunakan.',
    'Flow pemesanan lengkap untuk dine-in, pickup, delivery, dan reservasi.',
    'Integrasi payment gateway (Doku) untuk pembayaran digital.',
    'Database pelanggan dan transaksi dasar.',
    'Mekanisme human takeover dan eskalasi ke admin.',
    'Sinkronisasi menu otomatis dengan sistem POS.',
    'Dokumentasi alur sistem dan panduan operasional.',
    'Masa testing dan stabilisasi awal (1 minggu pasca go-live).',
]
for d in deliverables:
    add_numbered(d)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. RISIKO DAN MITIGASI
# ══════════════════════════════════════════════════════════════
add_heading_styled('13. Risiko dan Mitigasi', level=1)

# Risk table
rt = doc.add_table(rows=6, cols=3)
rt.style = 'Light Grid Accent 1'
rt.alignment = WD_TABLE_ALIGNMENT.CENTER

rt_headers = ['No.', 'Risiko', 'Mitigasi']
for i, h in enumerate(rt_headers):
    cell = rt.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

risks = [
    ('1', 'Kesalahan input dari pelanggan',
     'Sistem menyediakan pertanyaan bertahap dan ringkasan konfirmasi sebelum pembayaran.'),
    ('2', 'Kebutuhan penanganan manual',
     'Tersedia fitur eskalasi/human takeover ke admin untuk kasus di luar skenario otomatis.'),
    ('3', 'Gangguan server atau payment gateway',
     'Monitoring berkala, fallback operasional manual, dan pencatatan status order.'),
    ('4', 'Perubahan menu dan harga',
     'Data menu dirancang agar dapat diperbarui tanpa perubahan besar pada sistem inti. Sinkronisasi otomatis dengan POS.'),
    ('5', 'Adaptasi tim operasional',
     'Pendampingan awal, training singkat untuk staf, dan penyederhanaan alur kerja.'),
]
for i, (no, risk, mitigation) in enumerate(risks):
    row = rt.rows[i + 1]
    row.cells[0].text = no
    row.cells[1].text = risk
    row.cells[2].text = mitigation
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 14. PENUTUP
# ══════════════════════════════════════════════════════════════
add_heading_styled('14. Penutup', level=1)
add_para(
    'Implementasi WhatsApp AI Chatbot merupakan langkah yang relevan dan strategis untuk '
    'mendukung transformasi operasional Kedai Ngupi-Ngupi ke arah yang lebih efisien, '
    'terintegrasi, dan modern. Sistem ini diharapkan dapat membantu kedai dalam meningkatkan '
    'kualitas pelayanan, mempercepat alur transaksi, mengurangi beban kerja manual, serta '
    'membangun kanal pemesanan langsung yang lebih sehat secara bisnis.'
)
add_para(
    'Dengan memanfaatkan WhatsApp sebagai kanal utama, Kedai Ngupi-Ngupi dapat mendekatkan '
    'layanan ke pelanggan melalui media yang sudah akrab digunakan sehari-hari. Pada saat yang '
    'sama, kedai memiliki peluang untuk menekan ketergantungan terhadap platform pihak ketiga, '
    'menjaga margin keuntungan, dan membangun pengalaman transaksi yang lebih cepat dan lebih '
    'personal bagi setiap pelanggan.'
)
add_para(
    'Dalam jangka menengah, sistem ini juga dapat menjadi fondasi untuk pengembangan program '
    'loyalitas, promosi yang lebih terarah, dan pengelolaan data pelanggan yang lebih baik. '
    'Proposal ini disusun sebagai bahan pertimbangan dan dasar diskusi lebih lanjut terkait '
    'implementasi sistem.'
)
add_para(
    'Atas perhatian dan kerja samanya, kami ucapkan terima kasih.'
)

add_blank(3)

# Signature
p = doc.add_paragraph()
run = p.add_run('Hormat kami,')
run.font.size = Pt(11)

add_blank(2)

p = doc.add_paragraph()
run = p.add_run('Muhammad Rasyid Ridho')
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run('Developer WhatsApp AI Chatbot')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── Save ──
out_path = '/home/ubuntu/workspace-sobatngupi/PROPOSAL_IMPLEMENTASI_WHATSAPP_AI_CHATBOT_REVISI.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
print(f'Size: {os.path.getsize(out_path)} bytes')
